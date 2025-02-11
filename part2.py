# # file: app_rag.py
# import os
# from typing import List
# import streamlit as st
# from langchain_community.document_loaders import PDFPlumberLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain_community.llms import OpenAI
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()
# openai_api_key = os.getenv('OPENAI_API_KEY')
# os.environ["OPENAI_API_KEY"] = openai_api_key

# # Define HanaBOT class
# class HanaBOT:
#     def __init__(self, index_path="faiss_index"):
#         self.embeddings = OpenAIEmbeddings()
#         self.llm = OpenAI(temperature=0)
#         self.vectorstore = None
#         self.index_path = index_path
        
#     def load_pdf(self, pdf_path: str) -> List[str]:
#         loader = PDFPlumberLoader(pdf_path)
#         pages = loader.load()
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=1000,
#             chunk_overlap=200
#         )
#         texts = text_splitter.split_documents(pages)
#         return texts
    
#     def process_and_store(self, texts: List[str]):
#         self.vectorstore = FAISS.from_documents(texts, self.embeddings)
#         self.vectorstore.save_local(self.index_path)
        
#     def load_existing_index(self):
#         if os.path.exists(self.index_path):
#             self.vectorstore = FAISS.load_local(self.index_path, self.embeddings)
        
#     def retrieve_relevant_docs(self, query: str, k: int = 5) -> List[str]:
#         if not self.vectorstore:
#             self.load_existing_index()
#         retriever = self.vectorstore.as_retriever(search_kwargs={"k": k})
#         docs = retriever.get_relevant_documents(query)
#         return [doc.page_content for doc in docs]
    
#     def generate_answer(self, query: str, context: List[str]) -> str:
#         # Combine context for LLM input
#         context_text = "\n".join(context)
#         prompt = f"""
#         You are a helpful AI assistant. Use the provided  context_text to answer the question accurately with some description regarding the answer.
#         If the context contains  the  synonymn words  for the given query, please use them to answer the question.
#         If the context does not contain the answer, say "I don't know" instead of making up information.

#         ### Context:
#         {context_text}
        
#         ### Question:
#         {query}

#         ### Answer:
#         """
#         return self.llm(prompt)

# # Streamlit App with RAG
# def main():
#     st.title("HANA_BOT")
    
#     # File uploader
#     uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    
#     # Textbox for user query
#     user_query = st.text_input("Enter your query:")
    
#     # Button to process the query
#     if st.button("Submit"):
#         if uploaded_file is None:
#             st.error("Please upload a PDF file.")
#             return
        
#         if not user_query.strip():
#             st.error("Please enter a query.")
#             return

#         # Process the uploaded file and query
#         try:
#             # Save the uploaded file temporarily
#             temp_pdf_path = f"temp_{uploaded_file.name}"
#             with open(temp_pdf_path, "wb") as temp_file:
#                 temp_file.write(uploaded_file.read())
            
#             # Initialize HanaBOT
#             bot = HanaBOT(index_path="faiss_index")
#             texts = bot.load_pdf(temp_pdf_path)
#             bot.process_and_store(texts)
            
#             # Retrieve relevant docs and generate answer
#             relevant_docs = bot.retrieve_relevant_docs(user_query, k=5)
#             answer = bot.generate_answer(user_query, relevant_docs)
            
#             # Display results
#             st.subheader("Retrieved Context:")
#             for idx, doc in enumerate(relevant_docs):
#                 st.write(f"**Doc {idx + 1}:** {doc}")
#             st.subheader("Answer:")
#             st.success(answer)
            
#             # Clean up temporary file
#             os.remove(temp_pdf_path)
#         except Exception as e:
#             st.error(f"An error occurred: {str(e)}")

# if __name__ == "__main__":
#     main()




#Pressnt usage code

#Langchain Non agent code
# # file: app_rag.py
# import os
# from typing import List
# import streamlit as st
# from docx import Document  # For processing .docx files
# from langchain.schema import Document as LangChainDocument  # Import Document schema
# from langchain_community.document_loaders import PDFPlumberLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain_openai import ChatOpenAI  # GPT-4o Mini
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()
# openai_api_key = os.getenv('OPENAI_API_KEY')
# os.environ["OPENAI_API_KEY"] = openai_api_key

# # Define HanaBOT class
# class HanaBOT:
#     def __init__(self, index_path="faiss_index"):
#         self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
#         self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)  
#         self.index_path = index_path
#         self.vectorstore = None
#         self.load_existing_index()
        
#     def load_file(self, file_path: str, file_type: str) -> List[LangChainDocument]:
#         """Loads and splits file content into chunks based on file type"""
#         if file_type == "pdf":
#             # Use PDFPlumberLoader for PDF files
#             loader = PDFPlumberLoader(file_path)
#             pages = loader.load()
#         elif file_type == "docx":
#             # Use python-docx for DOCX files
#             doc = Document(file_path)
#             # Convert paragraphs into LangChainDocument objects
#             pages = [LangChainDocument(page_content=para.text) for para in doc.paragraphs if para.text.strip()]
#         else:
#             raise ValueError("Unsupported file type")

#         # Split the documents into chunks
#         text_splitter = RecursiveCharacterTextSplitter(
#             chunk_size=1000,
#             chunk_overlap=200
#         )
#         texts = text_splitter.split_documents(pages)
#         return texts

#     def process_and_store(self, texts: List[LangChainDocument]):
#         """Creates FAISS index only once and stores embeddings"""
#         self.vectorstore = FAISS.from_documents(texts, self.embeddings)
#         self.vectorstore.save_local(self.index_path)
    
#     def load_existing_index(self):
#         """Loads existing FAISS index if available"""
#         if os.path.exists(self.index_path):
#             self.vectorstore = FAISS.load_local(
#                 self.index_path, 
#                 self.embeddings, 
#                 allow_dangerous_deserialization=True
#             )
    
#     def retrieve_relevant_docs(self, query: str, k: int = 5) -> List[str]:
#         """Retrieves the most relevant documents based on the query"""
#         if not self.vectorstore:
#             st.error("No document has been processed. Please upload and process a file first.")
#             return []
        
#         retriever = self.vectorstore.as_retriever(search_kwargs={"k": k})
#         docs = retriever.get_relevant_documents(query)
#         return [doc.page_content for doc in docs]
    
#     def generate_answer(self, query: str, context: List[str]) -> str:
#         """Uses retrieved context to generate an answer"""
#         context_text = "\n".join(context)
#         prompt = f"""
#         Use the provided context to answer the user questions. The entire context may not be related to user question, so answer wisely from the context.
#         If the answer is not available in the context, please respond with "I couldn't find relevant information about that in the provided documents.
#         If the context contains synonym words for the given query, please use them in the answer.
#         Just give the answer to the question as it is  in the document or pdf without any additional information or summarising the information.


#         ### Context:
#         {context_text}
        
#         ### Question:
#         {query}

#         """
#         return self.llm.invoke(prompt).content

# # Streamlit App with RAG
# def main():
#     st.title("HANA_BOT")

#     # Initialize HanaBOT
#     bot = HanaBOT(index_path="faiss_index")

#     # File uploader
#     uploaded_file = st.file_uploader("Select a file (PDF or DOCX)", type=["pdf", "docx"])

#     # Process button (only runs once)
#     if uploaded_file and st.button("Process & Generate Embeddings"):
#         # Save the uploaded file temporarily
#         temp_file_path = f"temp_{uploaded_file.name}"
#         with open(temp_file_path, "wb") as temp_file:
#             temp_file.write(uploaded_file.read())

#         # Detect file type and process it
#         file_extension = uploaded_file.name.split(".")[-1].lower()

#         try:
#             texts = bot.load_file(temp_file_path, file_extension)
#             bot.process_and_store(texts)
#             st.success(f"{file_extension.upper()} file processed successfully! Embeddings have been generated.")
#         except Exception as e:
#             st.error(f"Error processing file: {str(e)}")
#         finally:
#             os.remove(temp_file_path)

#     # Query input
#     user_query = st.text_input("Enter your query:")

#     if st.button("Submit Query"):
#         if not user_query.strip():
#             st.error("Please enter a query.")
#             return

#         try:
#             # Retrieve relevant docs and generate answer
#             relevant_docs = bot.retrieve_relevant_docs(user_query, k=5)
#             answer = bot.generate_answer(user_query, relevant_docs)

#             # Display results
#             st.subheader("Retrieved Context:")
#             for idx, doc in enumerate(relevant_docs):
#                 st.write(f"**Doc {idx + 1}:** {doc}")

#             st.subheader("Answer:")
#             st.success(answer)
#         except Exception as e:
#             st.error(f"An error occurred: {str(e)}")

# if __name__ == "__main__":
#     main()


# import os
# import pandas as pd
# import streamlit as st
# from typing import List
# from sqlalchemy import create_engine, text, inspect
# import re, io, logging
# from docx import Document
# from langchain.schema import Document as LangChainDocument
# from langchain_community.document_loaders import PDFPlumberLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain_openai import ChatOpenAI
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()
# openai_api_key = os.getenv('OPENAI_API_KEY')
# os.environ["OPENAI_API_KEY"] = openai_api_key

# # Configure logging
# logger = logging.getLogger(__name__)
# logging.basicConfig(level=logging.INFO)

# class DatabaseManager:
#     def __init__(self, database_url: str):
#         self.database_url = database_url
#         self.engine = create_engine(database_url)
#         self.table_name = "new_excel_data"

#     def store_dataframe(self, file_path: str, if_exists: str = 'append') -> None:
#         try:
#             df = pd.read_excel(file_path) if file_path.endswith(".xlsx") else pd.read_csv(file_path)
#             df.to_sql(self.table_name, self.engine, if_exists=if_exists, index=False)
#             logger.info(f"Data stored successfully in {self.table_name}")
#         except Exception as e:
#             logger.error(f"Error storing DataFrame: {str(e)}")
#             raise

#     def execute_query(self, query: str) -> List[str]:
#         try:
#             with self.engine.connect() as connection:
#                 result_set = connection.execute(text(query))
#                 return [str(row) for row in result_set]
#         except Exception as e:
#             logger.error(f"Error executing query: {str(e)}")
#             return [str(e)]

# class HanaBOT:
#     def __init__(self, index_path="faiss_index"):
#         self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
#         self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
#         self.index_path = index_path
#         self.vectorstore = None
#         self.load_existing_index()
    
#     def parse_natural_query(self, user_query: str) -> str:
#         """Converts natural language query into SQL query using LLM if related to database"""
#         if "dataset" in user_query.lower() or "table" in user_query.lower():
#             prompt = f"Convert the following natural language query into a valid SQL query for a table named 'new_excel_data': \n{user_query}"
#             return self.llm.invoke(prompt).content.strip()
#         return None
    
#     def load_file(self, file_path: str, file_type: str) -> List[LangChainDocument]:
#         if file_type == "pdf":
#             loader = PDFPlumberLoader(file_path)
#             pages = loader.load()
#         elif file_type == "docx":
#             doc = Document(file_path)
#             pages = [LangChainDocument(page_content=para.text) for para in doc.paragraphs if para.text.strip()]
#         else:
#             raise ValueError("Unsupported file type")

#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#         texts = text_splitter.split_documents(pages)
#         return texts
    
#     def process_and_store(self, texts: List[LangChainDocument]):
#         self.vectorstore = FAISS.from_documents(texts, self.embeddings)
#         self.vectorstore.save_local(self.index_path)
    
#     def load_existing_index(self):
#         if os.path.exists(self.index_path):
#             self.vectorstore = FAISS.load_local(self.index_path, self.embeddings, allow_dangerous_deserialization=True)
    
#     def retrieve_relevant_docs(self, query: str, k: int = 5) -> List[str]:
#         if not self.vectorstore:
#             st.error("No document has been processed. Please upload and process a file first.")
#             return []
#         retriever = self.vectorstore.as_retriever(search_kwargs={"k": k})
#         docs = retriever.get_relevant_documents(query)
#         return [doc.page_content for doc in docs]
    
#     def generate_answer(self, query: str, context: List[str]) -> str:
#         context_text = "\n".join(context)
#         prompt = f"""
#         Use the provided context to answer the user questions. The entire context may not be related to user question, so answer wisely from the context.
#         \nIf the answer is not available in the context, please respond with "I couldn't find relevant information about that in the provided documents."

#         ### Context:
#         {context_text}
#         ### Question:
#         {query}
#         """
#         return self.llm.invoke(prompt).content

# def main():
#     st.title("HANA_BOT")

#     bot = HanaBOT(index_path="faiss_index")
#     db_manager = DatabaseManager(database_url="postgresql://test_owner:tcWI7unQ6REA@ep-yellow-recipe-a5fny139.us-east-2.aws.neon.tech:5432/test")

#     uploaded_file = st.file_uploader("Select a file (PDF, DOCX, Excel, CSV)", type=["pdf", "docx", "xlsx", "csv"])
    
#     if uploaded_file and st.button("Process File"):
#         temp_file_path = f"temp_{uploaded_file.name}"
#         with open(temp_file_path, "wb") as temp_file:
#             temp_file.write(uploaded_file.read())

#         file_extension = uploaded_file.name.split(".")[-1].lower()

#         try:
#             if file_extension in ["xlsx", "csv"]:
#                 db_manager.store_dataframe(temp_file_path)
#                 st.success(f"{file_extension.upper()} file stored in database successfully!")
#             else:
#                 texts = bot.load_file(temp_file_path, file_extension)
#                 bot.process_and_store(texts)
#                 st.success(f"{file_extension.upper()} file processed successfully! Embeddings generated.")
#         except Exception as e:
#             st.error(f"Error processing file: {str(e)}")
#         finally:
#             os.remove(temp_file_path)
    
#     user_query = st.text_input("Enter your query:")
    
#     if st.button("Submit Query"):
#         if not user_query.strip():
#             st.error("Please enter a query.")
#             return

#         try:
#             sql_query = bot.parse_natural_query(user_query)
#             if sql_query:
#                 results = db_manager.execute_query(sql_query)
#                 st.subheader("Query Results:")
#                 for result in results:
#                     st.write(result)
#             else:
#                 relevant_docs = bot.retrieve_relevant_docs(user_query, k=5)
#                 answer = bot.generate_answer(user_query, relevant_docs)
#                 st.subheader("Retrieved Context:")
#                 for idx, doc in enumerate(relevant_docs):
#                     st.write(f"**Doc {idx + 1}:** {doc}")
#                 st.subheader("Answer:")
#                 st.success(answer)
#         except Exception as e:
#             st.error(f"An error occurred: {str(e)}")

# if __name__ == "__main__":
#     main()



# # Code with doc,pdf,normal image.
# import os
# import pytesseract
# from PIL import Image
# from typing import List
# import streamlit as st
# from docx import Document
# from langchain.schema import Document as LangChainDocument
# from langchain_community.document_loaders import PDFPlumberLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_openai import OpenAIEmbeddings
# from langchain_community.vectorstores import FAISS
# from langchain_openai import ChatOpenAI
# from dotenv import load_dotenv

# # Load environment variables
# load_dotenv()
# openai_api_key = os.getenv('OPENAI_API_KEY')
# os.environ["OPENAI_API_KEY"] = openai_api_key

# # Define HanaBOT class
# class HanaBOT:
#     def __init__(self, index_path="faiss_index"):
#         self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
#         self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
#         self.index_path = index_path
#         self.vectorstore = None
#         self.load_existing_index()

#     def load_file(self, file_path: str, file_type: str) -> List[LangChainDocument]:
#         """Loads and splits file content into chunks based on file type"""
#         if file_type == "pdf":
#             loader = PDFPlumberLoader(file_path)
#             pages = loader.load()
#         elif file_type == "docx":
#             doc = Document(file_path)
#             pages = [LangChainDocument(page_content=para.text) for para in doc.paragraphs if para.text.strip()]
#         elif file_type in ["png", "jpg", "jpeg"]:
#             image = Image.open(file_path)
#             extracted_text = pytesseract.image_to_string(image)
#             pages = [LangChainDocument(page_content=extracted_text)]
#         else:
#             raise ValueError("Unsupported file type")

#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
#         texts = text_splitter.split_documents(pages)
#         return texts

#     def process_and_store(self, texts: List[LangChainDocument]):
#         """Creates FAISS index only once and stores embeddings"""
#         self.vectorstore = FAISS.from_documents(texts, self.embeddings)
#         self.vectorstore.save_local(self.index_path)
    
#     def load_existing_index(self):
#         """Loads existing FAISS index if available"""
#         if os.path.exists(self.index_path):
#             self.vectorstore = FAISS.load_local(
#                 self.index_path, 
#                 self.embeddings, 
#                 allow_dangerous_deserialization=True
#             )
    
#     def retrieve_relevant_docs(self, query: str, k: int = 5) -> List[str]:
#         """Retrieves the most relevant documents based on the query"""
#         if not self.vectorstore:
#             st.error("No document has been processed. Please upload and process a file first.")
#             return []
        
#         retriever = self.vectorstore.as_retriever(search_kwargs={"k": k})
#         docs = retriever.get_relevant_documents(query)
#         return [doc.page_content for doc in docs]
    
#     def generate_answer(self, query: str, context: List[str]) -> str:
#         """Uses retrieved context to generate an answer"""
#         context_text = "\n".join(context)
#         prompt = f"""
#           Use the provided context to answer the user questions. The entire context may not be related to user question, so answer wisely from the context.
#           If the answer is not available in the context, please respond with "I couldn't find relevant information about that in the provided documents."

        
        
#         ### Context:
#         {context_text}
        
#         ### Question:
#         {query}
#         """
#         return self.llm.invoke(prompt).content

# # Streamlit App with RAG
# def main():
#     st.title("HANA_BOT with Image Support")

#     # Initialize HanaBOT
#     bot = HanaBOT(index_path="faiss_index")

#     # File uploader
#     uploaded_file = st.file_uploader("Select a file (PDF, DOCX, PNG, JPG, JPEG)", type=["pdf", "docx", "png", "jpg", "jpeg"])

#     # Process button (only runs once)
#     if uploaded_file and st.button("Process & Generate Embeddings"):
#         temp_file_path = f"temp_{uploaded_file.name}"
#         with open(temp_file_path, "wb") as temp_file:
#             temp_file.write(uploaded_file.read())

#         file_extension = uploaded_file.name.split(".")[-1].lower()

#         try:
#             texts = bot.load_file(temp_file_path, file_extension)
#             bot.process_and_store(texts)
#             st.success(f"{file_extension.upper()} file processed successfully! Embeddings generated.")
#         except Exception as e:
#             st.error(f"Error processing file: {str(e)}")
#         finally:
#             os.remove(temp_file_path)

#     # Query input
#     user_query = st.text_input("Enter your query:")

#     if st.button("Submit Query"):
#         if not user_query.strip():
#             st.error("Please enter a query.")
#             return

#         try:
#             relevant_docs = bot.retrieve_relevant_docs(user_query, k=5)
#             answer = bot.generate_answer(user_query, relevant_docs)

#             st.subheader("Retrieved Context:")
#             for idx, doc in enumerate(relevant_docs):
#                 st.write(f"**Doc {idx + 1}:** {doc}")

#             st.subheader("Answer:")
#             st.success(answer)
#         except Exception as e:
#             st.error(f"An error occurred: {str(e)}")

# if __name__ == "__main__":
#     main()





import os
import pytesseract
import cv2
import numpy as np
from PIL import Image
from typing import List
import streamlit as st
from docx import Document
from langchain.schema import Document as LangChainDocument
from langchain_community.document_loaders import PDFPlumberLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
openai_api_key = os.getenv('OPENAI_API_KEY')
os.environ["OPENAI_API_KEY"] = openai_api_key

# Define HanaBOT class
class HanaBOT:
    def __init__(self, index_path="faiss_index"):
        self.embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
        self.llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0)
        self.index_path = index_path
        self.vectorstore = None
        self.load_existing_index()

    def preprocess_image(self, image_path: str) -> Image:
        """Preprocess the image for better OCR accuracy (for both printed and handwritten text)."""
        image = cv2.imread(image_path)

        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

        # Apply adaptive thresholding to enhance text visibility
        processed_image = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )

        # Apply median blur to reduce noise
        processed_image = cv2.medianBlur(processed_image, 3)

        # Save temp processed image
        temp_processed_path = "processed_temp.png"
        cv2.imwrite(temp_processed_path, processed_image)

        return temp_processed_path

    def extract_text_from_image(self, image_path: str) -> str:
        """Extracts text from both printed and handwritten images using OCR."""
        processed_image_path = self.preprocess_image(image_path)
        extracted_text = pytesseract.image_to_string(
            Image.open(processed_image_path),
            config="--oem 1 --psm 6"  # LSTM OCR with automatic segmentation
        )
        os.remove(processed_image_path)  # Cleanup temp processed file
        return extracted_text.strip()

    def load_file(self, file_path: str, file_type: str) -> List[LangChainDocument]:
        """Loads and splits file content into chunks based on file type."""
        if file_type == "pdf":
            loader = PDFPlumberLoader(file_path)
            pages = loader.load()
        elif file_type == "docx":
            doc = Document(file_path)
            pages = [LangChainDocument(page_content=para.text) for para in doc.paragraphs if para.text.strip()]
        elif file_type in ["png", "jpg", "jpeg"]:
            extracted_text = self.extract_text_from_image(file_path)
            pages = [LangChainDocument(page_content=extracted_text)]
        else:
            raise ValueError("Unsupported file type")

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        texts = text_splitter.split_documents(pages)
        return texts

    def process_and_store(self, texts: List[LangChainDocument]):
        """Creates FAISS index only once and stores embeddings."""
        self.vectorstore = FAISS.from_documents(texts, self.embeddings)
        self.vectorstore.save_local(self.index_path)
    
    def load_existing_index(self):
        """Loads existing FAISS index if available."""
        if os.path.exists(self.index_path):
            self.vectorstore = FAISS.load_local(
                self.index_path, 
                self.embeddings, 
                allow_dangerous_deserialization=True
            )
    
    def retrieve_relevant_docs(self, query: str, k: int = 5) -> List[str]:
        """Retrieves the most relevant documents based on the query."""
        if not self.vectorstore:
            st.error("No document has been processed. Please upload and process a file first.")
            return []
        
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": k})
        docs = retriever.get_relevant_documents(query)
        return [doc.page_content for doc in docs]
    
    def generate_answer(self, query: str, context: List[str]) -> str:
        """Uses retrieved context to generate an answer."""
        context_text = "\n".join(context)
        prompt = f"""
          Use the provided context to answer the user questions. The entire context may not be related to user question, so answer wisely from the context.
          If the answer is not available in the context, please respond with "I couldn't find relevant information about that in the provided documents."

          You have to give the information whatever present in the document,pdf and image without any additional information or summarising the information.

        ### Context:
        {context_text}
        
        ### Question:
        {query}
        """
        return self.llm.invoke(prompt).content

# Streamlit App with RAG
def main():
    st.title("Chat2Doc") 

    # Initialize HanaBOT
    bot = HanaBOT(index_path="faiss_index")

    # File uploader
    uploaded_file = st.file_uploader("Select a file (PDF, DOCX, PNG, JPG, JPEG)", type=["pdf", "docx", "png", "jpg", "jpeg"])

    # Process button (only runs once)
    if uploaded_file and st.button("Process & Generate Embeddings"):
        temp_file_path = f"temp_{uploaded_file.name}"
        with open(temp_file_path, "wb") as temp_file:
            temp_file.write(uploaded_file.read())

        file_extension = uploaded_file.name.split(".")[-1].lower()

        try:
            texts = bot.load_file(temp_file_path, file_extension)
            bot.process_and_store(texts)
            st.success(f"{file_extension.upper()} file processed successfully! Embeddings generated.")
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
        finally:
            os.remove(temp_file_path)

    # Query input
    user_query = st.text_input("Enter your query:")

    if st.button("Submit Query"):
        if not user_query.strip():
            st.error("Please enter a query.")
            return

        try:
            relevant_docs = bot.retrieve_relevant_docs(user_query, k=5)
            answer = bot.generate_answer(user_query, relevant_docs)

            st.subheader("Retrieved Context:")
            for idx, doc in enumerate(relevant_docs):
                st.write(f"**Doc {idx + 1}:** {doc}")

            st.subheader("Answer:")
            st.success(answer)
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
